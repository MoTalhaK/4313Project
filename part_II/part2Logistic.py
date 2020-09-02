import pandas as pd
import docx
from sklearn.linear_model import LogisticRegression
from sklearn.metrics import confusion_matrix

f1_list = []  # list for f1-scores
p_list = []  # list for precision
r_list = []  # list for recall

tp_list = []
fp_list = []
fn_list = []
tn_list = []


def learn_dt(file_train, file_test, penalty, dual, tol,
             C, fit_intercept, intercept_scaling, class_weight,
             random_state, solver, max_iter,
             verbose, warm_start, n_jobs, l1_ratio):
    # load the training data as a matrix
    dataset = pd.read_csv(file_train, header=0)

    # separate the data from the target attributes
    train_data = dataset.drop('500_Buggy?', axis=1)

    # remove unnecessary features
    train_data = train_data.drop('change_id', axis=1)
    train_data = train_data.drop('412_full_path', axis=1)
    train_data = train_data.drop('411_commit_time', axis=1)

    # the lables of training data. `label` is the title of the  last column in your CSV files
    train_target = dataset.iloc[:, -1]

    # load the testing data
    dataset2 = pd.read_csv(file_test, header=0)

    # separate the data from the target attributes
    test_data = dataset2.drop('500_Buggy?', axis=1)

    # remove unnecessary features
    test_data = test_data.drop('change_id', axis=1)
    test_data = test_data.drop('412_full_path', axis=1)
    test_data = test_data.drop('411_commit_time', axis=1)

    # the lables of test data
    test_target = dataset2.iloc[:, -1]

    l_regression = LogisticRegression(penalty=penalty, dual=dual, tol=tol,
                                      C=C, fit_intercept=fit_intercept, intercept_scaling=intercept_scaling,
                                      class_weight=class_weight,
                                      random_state=random_state, solver=solver, max_iter=max_iter,
                                      verbose=verbose, warm_start=warm_start, n_jobs=n_jobs, l1_ratio=l1_ratio)

    test_pred = l_regression.fit(train_data, train_target).predict(test_data)
    conf_matrix = confusion_matrix(test_target, test_pred, labels=[0, 1])
    TP = conf_matrix[0][0]
    FP = conf_matrix[0][1]
    FN = conf_matrix[1][0]
    TN = conf_matrix[1][1]
    tp_list.append(TP)
    fp_list.append(FP)
    fn_list.append(FN)
    tn_list.append(TN)


def param_test_jr(c_val):
    for i in range(0, 6):
        learn_dt("../data/jackrabbit/" + str(i) + "/train.csv", "../data/jackrabbit/" + str(i) + "/test.csv",
                 penalty='l2', dual=False, tol=0.0001, C=c_val, fit_intercept=True, intercept_scaling=1,
                 class_weight=None, random_state=30, solver='liblinear', max_iter=10000, verbose=0,
                 warm_start=False, n_jobs=None, l1_ratio=None)

    total_tp = sum(tp_list)
    total_fp = sum(fp_list)
    total_fn = sum(fn_list)
    total_tn = sum(tn_list)

    precision = total_tp / (total_tp + total_fp)
    recall = total_tp / (total_tp + total_fn)
    F1 = (2 * precision * recall) / (precision + recall)
    print("...")

    f1_list.append(round(F1, 2))
    p_list.append(round(precision, 2))
    r_list.append(round(recall, 2))

    tp_list.clear()
    fp_list.clear()
    fn_list.clear()
    tn_list.clear()


def param_test_jdt(c_val):
    # jdt
    for i in range(0, 6):
        learn_dt("../data/jdt/" + str(i) + "/train.csv", "../data/jdt/" + str(i) + "/test.csv",
                 penalty='l2', dual=False, tol=0.0001, C=c_val, fit_intercept=True, intercept_scaling=1,
                 class_weight=None, random_state=30, solver='liblinear', max_iter=10000, verbose=0,
                 warm_start=False, n_jobs=None, l1_ratio=None)

    total_tp = sum(tp_list)
    total_fp = sum(fp_list)
    total_fn = sum(fn_list)
    total_tn = sum(tn_list)

    precision = total_tp / (total_tp + total_fp)
    recall = total_tp / (total_tp + total_fn)
    F1 = (2 * precision * recall) / (precision + recall)
    print("...")

    f1_list.append(round(F1, 2))
    p_list.append(round(precision, 2))
    r_list.append(round(recall, 2))

    tp_list.clear()
    fp_list.clear()
    fn_list.clear()
    tn_list.clear()


def param_test_lucene(c_val):
    # lucene
    for i in range(0, 6):
        learn_dt("../data/lucene/" + str(i) + "/train.csv", "../data/lucene/" + str(i) + "/test.csv",
                 penalty='l2', dual=False, tol=0.0001, C=c_val, fit_intercept=True, intercept_scaling=1,
                 class_weight=None, random_state=30, solver='liblinear', max_iter=10000, verbose=0,
                 warm_start=False, n_jobs=None, l1_ratio=None)

    total_tp = sum(tp_list)
    total_fp = sum(fp_list)
    total_fn = sum(fn_list)
    total_tn = sum(tn_list)

    precision = total_tp / (total_tp + total_fp)
    recall = total_tp / (total_tp + total_fn)
    F1 = (2 * precision * recall) / (precision + recall)
    print("...")

    f1_list.append(round(F1, 2))
    p_list.append(round(precision, 2))
    r_list.append(round(recall, 2))

    tp_list.clear()
    fp_list.clear()
    fn_list.clear()
    tn_list.clear()


def param_test_xorg(c_val):
    # xorg
    for i in range(0, 6):
        learn_dt("../data/xorg/" + str(i) + "/train.csv", "../data/xorg/" + str(i) + "/test.csv",
                 penalty='l2', dual=False, tol=0.0001, C=c_val, fit_intercept=True, intercept_scaling=1,
                 class_weight=None, random_state=30, solver='liblinear', max_iter=10000, verbose=0,
                 warm_start=False, n_jobs=None, l1_ratio=None)

    total_tp = sum(tp_list)
    total_fp = sum(fp_list)
    total_fn = sum(fn_list)
    total_tn = sum(tn_list)

    precision = total_tp / (total_tp + total_fp)
    recall = total_tp / (total_tp + total_fn)
    F1 = (2 * precision * recall) / (precision + recall)
    print("...")

    f1_list.append(round(F1, 2))
    p_list.append(round(precision, 2))
    r_list.append(round(recall, 2))

    tp_list.clear()
    fp_list.clear()
    fn_list.clear()
    tn_list.clear()


doc = docx.Document("test.docx")


def to_doc(d_frame):
    # global doc
    t = doc.add_table(d_frame.shape[0] + 1, d_frame.shape[1])

    for j in range(df.shape[-1]):
        t.cell(0, j).text = df.columns[j]

    # add the rest of the data frame
    for i in range(df.shape[0]):
        for j in range(df.shape[-1]):
            t.cell(i + 1, j).text = str(df.values[i, j])
    doc.save("test.docx")


n_val = [0.0001, 0.001, 0.01, 0.1, 1, 10, 50, 100, 500, 1000]
doc.add_paragraph("Using: solver = 'liblinear' and penalty = 'l2'...")
print("jackrabbit...\n")
for n in n_val:
    param_test_jr(n)

df = pd.DataFrame(list(zip(n_val, p_list, r_list, f1_list)),
                  columns=["C", "Precision", "Recall", "F1-Score"])

to_doc(df)
print(df)
f1_list.clear()
p_list.clear()
r_list.clear()

print("jdt...\n")
for n in n_val:
    param_test_jdt(n)

df = pd.DataFrame(list(zip(n_val, p_list, r_list, f1_list)),
                  columns=["C", "Precision", "Recall", "F1-Score"])

to_doc(df)
print(df)
f1_list.clear()
p_list.clear()
r_list.clear()

print("lucene...\n")
for n in n_val:
    param_test_lucene(n)

df = pd.DataFrame(list(zip(n_val, p_list, r_list, f1_list)),
                  columns=["C", "Precision", "Recall", "F1-Score"])

to_doc(df)
print(df)
f1_list.clear()
p_list.clear()
r_list.clear()

print("xorg...\n")
for n in n_val:
    param_test_xorg(n)

df = pd.DataFrame(list(zip(n_val, p_list, r_list, f1_list)),
                  columns=["C", "Precision", "Recall", "F1-Score"])

to_doc(df)
print(df)
f1_list.clear()
p_list.clear()
r_list.clear()
